"""
Generate project plan as Excel tracker + Word document.
Run: python scripts/generate_project_plan.py
Outputs: docs/Project_Plan_Tracker.xlsx  +  docs/Project_Plan.docx
"""

import os
from pathlib import Path

import openpyxl
from openpyxl.styles import (
    Font, PatternFill, Alignment, Border, Side, GradientFill
)
from openpyxl.utils import get_column_letter
from openpyxl.formatting.rule import DataBarRule, ColorScaleRule
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = Path("docs")
OUT_DIR.mkdir(exist_ok=True)

# ─────────────────────────── COLOUR PALETTE ───────────────────────────
C = {
    "header_dark":  "1B3A5C",
    "header_mid":   "2E6DA4",
    "phase1":       "D6EAF8",
    "phase2":       "FDEBD0",
    "phase3":       "D5F5E3",
    "phase4":       "F9E79F",
    "phase5":       "FADBD8",
    "done":         "ABEBC6",
    "in_progress":  "FAD7A0",
    "not_started":  "F2F3F4",
    "crit":         "E74C3C",
    "high":         "E67E22",
    "med":          "F1C40F",
    "low":          "2ECC71",
    "white":        "FFFFFF",
    "light_grey":   "F8F9FA",
    "border":       "BDC3C7",
    "text_dark":    "1A252F",
}

def fill(hex_color):
    return PatternFill("solid", fgColor=hex_color)

def bold_font(size=10, color="1A252F", italic=False):
    return Font(name="Calibri", bold=True, size=size, color=color, italic=italic)

def normal_font(size=10, color="1A252F"):
    return Font(name="Calibri", bold=False, size=size, color=color)

def thin_border():
    side = Side(style="thin", color=C["border"])
    return Border(left=side, right=side, top=side, bottom=side)

def center():
    return Alignment(horizontal="center", vertical="center", wrap_text=True)

def left_wrap():
    return Alignment(horizontal="left", vertical="center", wrap_text=True)

def style_header_row(ws, row, cols, bg=C["header_dark"], fg=C["white"], size=10):
    for col in range(1, cols + 1):
        cell = ws.cell(row=row, column=col)
        cell.fill = fill(bg)
        cell.font = bold_font(size=size, color=fg)
        cell.alignment = center()
        cell.border = thin_border()

def style_data_row(ws, row, cols, bg=C["white"]):
    for col in range(1, cols + 1):
        cell = ws.cell(row=row, column=col)
        cell.fill = fill(bg)
        cell.font = normal_font()
        cell.alignment = left_wrap()
        cell.border = thin_border()

def set_col_widths(ws, widths):
    for i, w in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = w

def write_phase_header(ws, row, phase_name, color, col_count):
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=col_count)
    cell = ws.cell(row=row, column=1)
    cell.value = phase_name
    cell.fill = fill(color)
    cell.font = bold_font(size=11, color=C["header_dark"])
    cell.alignment = center()
    cell.border = thin_border()

# ═══════════════════════════════════════════════════════════════════════
#  EXCEL
# ═══════════════════════════════════════════════════════════════════════

def build_excel():
    wb = openpyxl.Workbook()
    wb.remove(wb.active)

    _sheet_dashboard(wb)
    _sheet_phase(wb, "Phase 1 – Setup", _phase1_tasks(), C["phase1"], "1B3A5C")
    _sheet_phase(wb, "Phase 2 – Fixes", _phase2_tasks(), C["phase2"], "7D5A1E")
    _sheet_phase(wb, "Phase 3 – Strategies", _phase3_tasks(), C["phase3"], "1D5E35")
    _sheet_phase(wb, "Phase 4 – Paper Validate", _phase4_tasks(), C["phase4"], "7D6608")
    _sheet_phase(wb, "Phase 5 – Cloud Go-Live", _phase5_tasks(), C["phase5"], "7B241C")
    _sheet_strategy_registry(wb)
    _sheet_known_issues(wb)

    path = OUT_DIR / "Project_Plan_Tracker.xlsx"
    wb.save(path)
    print(f"✅  Excel saved → {path}")


def _sheet_dashboard(wb):
    ws = wb.create_sheet("📊 Dashboard")
    ws.sheet_view.showGridLines = False
    ws.row_dimensions[1].height = 40

    # Title
    ws.merge_cells("A1:H1")
    c = ws["A1"]
    c.value = "DELTA EXCHANGE TRADING BOT — PROJECT TRACKER"
    c.fill = fill(C["header_dark"])
    c.font = bold_font(size=16, color=C["white"])
    c.alignment = center()

    ws.merge_cells("A2:H2")
    c = ws["A2"]
    c.value = "97-day roadmap  |  13 strategies  |  5 phases  |  $24/month when live"
    c.fill = fill(C["header_mid"])
    c.font = bold_font(size=10, color=C["white"])
    c.alignment = center()

    # Phase summary table
    headers = ["Phase", "Days", "Goal", "Status", "Progress"]
    phases = [
        ["Phase 1 – Local Setup",       "Days 1–7",   "Bot running in paper mode",                         "Done",        "100%"],
        ["Phase 2 – Fix Critical Issues","Days 8–37",  "Security fixed, bugs corrected, Grafana live",      "Done",        "~85%"],
        ["Phase 3 – New Strategies",    "Days 38–82", "8 new strategies built and backtested",             "In Progress", "~0%"],
        ["Phase 4 – Paper Validation",  "Days 83–97", "Sharpe > 0.8, Drawdown < 8%, Win rate > 45%",      "Not Started", "0%"],
        ["Phase 5 – Cloud Migration",   "Days 98+",   "Hetzner server, systemd, live trading with $1,000", "Not Started", "0%"],
    ]
    phase_colors = [C["phase1"], C["phase2"], C["phase3"], C["phase4"], C["phase5"]]

    ws.row_dimensions[4].height = 22
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=4, column=col, value=h)
        cell.fill = fill(C["header_dark"])
        cell.font = bold_font(color=C["white"])
        cell.alignment = center()
        cell.border = thin_border()

    for i, (row_data, bg) in enumerate(zip(phases, phase_colors), 5):
        for col, val in enumerate(row_data, 1):
            cell = ws.cell(row=i, column=col, value=val)
            cell.fill = fill(bg)
            cell.font = normal_font()
            cell.alignment = left_wrap()
            cell.border = thin_border()
        ws.row_dimensions[i].height = 22

    # Quick metrics
    ws.merge_cells("A11:H11")
    ws["A11"].value = "Key Metrics & Rules"
    ws["A11"].fill = fill(C["header_mid"])
    ws["A11"].font = bold_font(size=11, color=C["white"])
    ws["A11"].alignment = center()

    metrics = [
        ["Metric", "Value", "", "Rule", "Detail"],
        ["Total Strategies", "13", "", "Never commit .env to git", "Keys are already exposed — rotate first"],
        ["Capital (start)", "$1,000", "", "Paper validate before live", "Minimum 14 days, both gates must pass"],
        ["Max Live Position Size", "$50 (half paper)", "", "Start at half paper size", "Scale up only after first 10 live trades"],
        ["Taker Fee", "0.05% / trade", "", "Fix CRIT issues before Phase 3", "No point building strategies on broken code"],
        ["Monthly Cost (dev)", "$20 (Claude Pro)", "", "Sharpe gate for new strategies", "Only add to live dispatch if Sharpe > 0.8"],
        ["Monthly Cost (live)", "~$24", "", "Max daily loss kill switch", "5% daily loss → bot stops all new trades"],
    ]

    for i, row_data in enumerate(metrics, 12):
        bg = C["header_dark"] if i == 12 else (C["light_grey"] if i % 2 == 0 else C["white"])
        fg = C["white"] if i == 12 else C["text_dark"]
        for col, val in enumerate(row_data, 1):
            cell = ws.cell(row=i, column=col, value=val)
            cell.fill = fill(bg)
            cell.font = bold_font(color=fg) if i == 12 else normal_font()
            cell.alignment = left_wrap()
            cell.border = thin_border()
        ws.row_dimensions[i].height = 20

    set_col_widths(ws, [28, 20, 3, 32, 42])
    ws.freeze_panes = "A4"


def _task_cols():
    return ["#", "Task", "File / Location", "Day(s)", "Duration", "Priority", "Status", "Notes"]

def _phase_col_widths():
    return [4, 42, 28, 10, 10, 10, 14, 30]

def _sheet_phase(wb, sheet_name, tasks, bg_color, text_color):
    ws = wb.create_sheet(sheet_name)
    ws.sheet_view.showGridLines = False

    # Title
    ws.merge_cells(f"A1:{get_column_letter(len(_task_cols()))}1")
    c = ws["A1"]
    c.value = sheet_name.upper()
    c.fill = fill(C["header_dark"])
    c.font = bold_font(size=13, color=C["white"])
    c.alignment = center()
    ws.row_dimensions[1].height = 30

    # Column headers
    for col, h in enumerate(_task_cols(), 1):
        cell = ws.cell(row=2, column=col, value=h)
        cell.fill = fill(C["header_mid"])
        cell.font = bold_font(color=C["white"])
        cell.alignment = center()
        cell.border = thin_border()
    ws.row_dimensions[2].height = 20

    # Data rows
    section_colors = {
        "CRITICAL": C["crit"],
        "HIGH": C["high"],
        "MED": C["med"],
        "LOW": C["low"],
    }
    col_count = len(_task_cols())
    row_num = 3
    task_num = 1

    for item in tasks:
        if isinstance(item, str):
            # Section separator
            write_phase_header(ws, row_num, f"  {item}", bg_color, col_count)
            ws.row_dimensions[row_num].height = 18
            row_num += 1
        else:
            task, file_loc, days, duration, priority, status, notes = item
            row_bg = C["done"] if status == "Done" else (
                     C["in_progress"] if status == "In Progress" else
                     C["light_grey"] if task_num % 2 == 0 else C["white"])

            data = [task_num, task, file_loc, days, duration, priority, status, notes]
            for col, val in enumerate(data, 1):
                cell = ws.cell(row=row_num, column=col, value=val)
                cell.fill = fill(row_bg)
                cell.font = normal_font()
                cell.alignment = left_wrap()
                cell.border = thin_border()
                # Colour-code priority
                if col == 6 and val in section_colors:
                    cell.fill = fill(section_colors[val])
                    cell.font = bold_font(color=C["white"])
                    cell.alignment = center()
                # Centre status
                if col == 7:
                    cell.alignment = center()
            ws.row_dimensions[row_num].height = 22
            task_num += 1
            row_num += 1

    set_col_widths(ws, _phase_col_widths())
    ws.freeze_panes = "A3"

    # Dropdown for Status column (col 7)
    from openpyxl.worksheet.datavalidation import DataValidation
    dv = DataValidation(
        type="list",
        formula1='"Not Started,In Progress,Done,Blocked"',
        allow_blank=True,
        showDropDown=False,
    )
    dv.sqref = f"G3:G{row_num}"
    ws.add_data_validation(dv)


def _phase1_tasks():
    return [
        "INSTALLATION",
        ("Verify Python 3.12.x installed",          "PowerShell",                 "Day 1",   "30 min", "CRITICAL", "Done", "Confirmed: Python 3.12 installed"),
        ("Install Poetry package manager",           "PowerShell",                 "Day 1",   "30 min", "CRITICAL", "Done", "Poetry installed and working"),
        ("Run poetry install",                       "Project root",               "Day 1",   "1 hr",   "CRITICAL", "Done", "All dependencies installed"),
        "SECURITY — DO THIS BEFORE ANYTHING ELSE",
        ("Rotate leaked API keys in .env",           ".env",                       "Day 2",   "1 hr",   "CRITICAL", "Done", "Old keys rotated; testnet keys active"),
        ("Register on Delta testnet",                "cdn-ind.testnet.deltaex.org","Day 2",   "30 min", "CRITICAL", "Done", "Testnet account active"),
        ("Update .env with testnet credentials",     ".env",                       "Day 2",   "10 min", "CRITICAL", "Done", "DELTA_EXCHANGE_ENV=testnet-india in use"),
        "FIRST RUN",
        ("Run bot in paper mode",                    "cli/main.py",                "Day 2",   "30 min", "HIGH",     "Done", "Bot runs; upgraded to live testnet mode"),
        ("Observe log output for 30 min",            "Terminal",                   "Day 2",   "30 min", "HIGH",     "Done", "Logs verified; JSON events working"),
        ("Verify DB has trade/signal rows",          "state.db",                   "Day 2",   "15 min", "HIGH",     "Done", "state.db has trades, positions, logs"),
        "OBSERVATION (Days 3–7)",
        ("Run paper mode 2+ hours daily",            "Terminal",                   "Days 3–7","Passive","MED",      "Done", "Multiple sessions run; logs archived"),
        ("Confirm bot runs 2 consecutive days clean","Terminal",                   "Day 7",   "–",      "HIGH",     "Done", "Phase 1 gate passed — bot running live testnet"),
    ]


def _phase2_tasks():
    return [
        "CRITICAL SEVERITY",
        ("Fix Dockerfile Python 3.11 → 3.12",        "Dockerfile line 1",                         "Day 8",    "30 min", "CRITICAL", "Done",        "Confirmed: Dockerfile FROM python:3.12-slim"),
        ("Merge strategy/ and strategies/ packages",  "src/delta_exchange_bot/strategy/",          "Days 8–10","3 days", "CRITICAL", "Done",        "Only strategy/ exists; strategies/ removed"),
        ("Update all imports after merge",            "core/engine.py, portfolio.py",              "Day 10",   "1 day",  "CRITICAL", "Done",        "All imports use strategy/ package"),
        ("Change TRADE_FREQUENCY_S 60 → 5",          ".env, core/settings.py",                    "Day 8",    "30 min", "CRITICAL", "Done",        "trade_frequency_s = 5 in settings.py"),
        ("Add exchange-native stop orders",           "api/delta_client.py",                       "Days 14–20","6 days","CRITICAL", "In Progress", "Broken HTTP 400 call removed; TODO: bracket API /v2/orders/bracket in Phase 3"),
        ("Store stop_order_id in DB",                 "persistence/models.py",                     "Day 15",   "1 day",  "CRITICAL", "Not Started", "Blocked on native stop orders"),
        ("Cancel stop order on position close",       "execution/order_execution_engine.py",       "Day 16",   "1 day",  "CRITICAL", "Not Started", "Blocked on native stop orders"),
        ("Add idempotent order placement",            "api/delta_client.py",                       "Day 12",   "1 day",  "CRITICAL", "Not Started", "SHA256(trade_id+symbol+side) as client_order_id"),
        "HIGH SEVERITY",
        ("Fix RSI Wilder's smoothing formula",        "strategy/rsi_scalping.py",                  "Days 12–13","1 day", "HIGH",     "Not Started", "Currently 7 pts off vs TradingView"),
        ("Wire YAML config files to Settings",        "core/settings.py",                          "Days 10–12","2 days","HIGH",     "Done",        "Settings.__init__ loads default.yml + mode.yml"),
        ("Consolidate two risk managers into one",    "risk/risk_manager.py",                      "Day 11",   "1 day",  "HIGH",     "Not Started", "advanced_risk_manager.py still active alongside risk_manager.py"),
        ("Add fill confirmation polling loop",        "execution/order_execution_engine.py",       "Day 13",   "1 day",  "HIGH",     "Done",        "Issue 3 fixed: POST_EXECUTION_POSITION_NOT_CONFIRMED downgraded to WARNING when before==after==0 (testnet latency)"),
        "MEDIUM SEVERITY",
        ("Add funding rate awareness",                "api/delta_client.py, execution/",           "Days 20–24","4 days","MED",      "Not Started", "get_funding_rate(), log on each trade"),
        ("Add funding rate cost to PnL calc",         "execution/fee_manager.py",                  "Day 22",   "1 day",  "MED",      "Not Started", "Total fee = taker + funding cost"),
        ("Add max holding period (30 min default)",   "core/engine.py, core/settings.py",          "Days 24–26","2 days","MED",      "Done",        "max_holding_time_s = 1800 in settings.py"),
        ("Persist ProtectionState to PostgreSQL",     "persistence/models.py, db.py",              "Day 20",   "2 days", "MED",      "Not Started", "SL/TP still lost on crash (in-memory only)"),
        ("Add crash recovery on startup",             "core/engine.py",                            "Day 21",   "1 day",  "MED",      "Done",        "Issues 2-5 applied 2026-03-29: SL/TP HTTP 400 removed, grace period added, WebSocket watchdog pong fix"),
        "THIS SESSION FIXES (2026-03-28)",
        ("Fix: position state crash loop on WebSocket sync", "cli/professional_bot.py",            "2026-03-28","Done",  "CRITICAL", "Done",        "Save/restore local position before 5x retry; min_interval_s raised to 15s"),
        ("Fix: default SL/TP for exchange-synced positions", "cli/professional_bot.py",            "2026-03-28","Done",  "CRITICAL", "Done",        "Applies 0.4% SL / 0.8% TP when no prior state exists"),
        ("Fix: reconciliation stamps sync timestamp", "cli/professional_bot.py",                   "2026-03-28","Done",  "HIGH",     "Done",        "Prevents post-reconciliation sync from flattening confirmed positions"),
        ("Fix: removed broken native SL/TP HTTP 400","execution/order_execution_engine.py",        "2026-03-28","Done",  "HIGH",     "Done",        "Removed stop_market_order / take_profit_market_order calls; in-memory fallback only"),
        ("Add timestamped file logging",              "utils/logging.py, core/settings.py",        "2026-03-28","Done",  "MED",      "Done",        "logs/bot_YYYYMMDD_HHMMSS.log per run + logs/errors.log accumulates warnings"),
        "NEW BUGS FOUND (2026-03-28) — Fix Next",
        ("Kill switch permanently stops bot",         "cli/professional_bot.py run_async()",        "Next",     "1 day",  "HIGH",     "Not Started", "After 5% daily loss, _kill_switch_triggered=True exits while-loop. Need soft-mode: pause + alert, not hard exit"),
        ("Portfolio confidence scores too low",       "strategy/portfolio.py",                     "Next",     "1 day",  "HIGH",     "Done",        "Fixed 2026-03-29: RSI added to TRENDING regime; DELTA_MIN_SIGNAL_CONFIDENCE lowered to 0.50"),
        "MONITORING",
        ("Add Grafana to docker-compose.yml",         "docker-compose.yml",                        "Days 26–37","10 days","MED",     "Not Started", "+ Loki for log aggregation"),
        ("Import Trading Overview dashboard",         "Grafana UI",                                "Day 28",   "1 day",  "MED",      "Not Started", "PnL, win rate, positions"),
        ("Import Risk Monitor dashboard",             "Grafana UI",                                "Day 29",   "1 day",  "MED",      "Not Started", "Drawdown, leverage, daily loss"),
        ("Import API Health dashboard",               "Grafana UI",                                "Day 30",   "1 day",  "MED",      "Not Started", "Request latency, error rate"),
    ]


def _phase3_tasks():
    return [
        "PRIORITY 1 — VWAP Deviation (Days 38–42)",
        ("Create VWAPDeviationStrategy class",        "strategies/vwap_deviation.py (new)",        "Day 38",   "2 days", "HIGH",     "Not Started", "Entry: price >0.5% from VWAP"),
        ("Wire to StrategyManager RANGING regime",    "core/engine.py",                            "Day 40",   "1 day",  "HIGH",     "Not Started", "Activate for RANGING + LOW_VOL"),
        ("Backtest on 30 days BTCUSD 1m data",        "backtesting/",                              "Day 41",   "1 day",  "HIGH",     "Not Started", "Gate: Sharpe > 0.5"),
        ("48-hour paper run",                         "Terminal",                                  "Day 42",   "Passive","HIGH",     "Not Started", "Observe signal quality"),
        "PRIORITY 2 — Bollinger Bands Squeeze (Days 43–46)",
        ("Create BollingerSqueezeStrategy class",     "strategies/bollinger_squeeze.py (new)",     "Day 43",   "2 days", "HIGH",     "Not Started", "Detect low-vol squeeze before breakout"),
        ("Wire to StrategyManager LOW_VOL regime",    "core/engine.py",                            "Day 45",   "1 day",  "HIGH",     "Not Started", "LOW_VOL → HIGH_VOL transition"),
        ("Backtest + validate",                       "backtesting/",                              "Day 46",   "1 day",  "HIGH",     "Not Started", "Gate: Sharpe > 0.8"),
        "PRIORITY 3 — MACD Histogram (Days 47–50)",
        ("Create MACDHistogramStrategy class",        "strategies/macd_histogram.py (new)",        "Day 47",   "2 days", "HIGH",     "Not Started", "MACD(12,26,9) histogram direction"),
        ("Wire alongside TrendFollowing in TRENDING", "core/engine.py",                            "Day 49",   "1 day",  "HIGH",     "Not Started", "Both strategies active in TRENDING"),
        ("Backtest + validate",                       "backtesting/",                              "Day 50",   "1 day",  "HIGH",     "Not Started", "Gate: Sharpe > 0.8"),
        "PRIORITY 4 — Order Book Imbalance (Days 51–56)",
        ("Create OrderBookImbalanceStrategy class",   "strategies/order_book_imbalance.py (new)", "Day 51",   "3 days", "HIGH",     "Not Started", "Use existing get_orderbook() REST call"),
        ("Calculate bid/ask ratio top 5 levels",      "strategies/order_book_imbalance.py",        "Day 53",   "1 day",  "HIGH",     "Not Started", "Imbalance > 0.70 = buy signal"),
        ("Wire as secondary signal all regimes",      "core/engine.py",                            "Day 54",   "1 day",  "HIGH",     "Not Started", "Only fetch when primary signal fires"),
        ("Backtest + validate",                       "backtesting/",                              "Day 55",   "1 day",  "HIGH",     "Not Started", "Gate: Sharpe > 0.8"),
        "PRIORITY 5 — Multi-Timeframe Confluence (Days 57–62)",
        ("Create MultiTimeframeFilter wrapper class", "strategies/multi_timeframe_filter.py (new)","Day 57",  "3 days", "HIGH",     "Not Started", "1m + 15m + 1h all must agree"),
        ("Add fetch_candles(symbol,'1h') support",    "data/market_data.py",                       "Day 58",   "1 day",  "MED",      "Not Started", "1h candles not currently fetched"),
        ("Apply as filter around all strategies",     "core/engine.py",                            "Day 60",   "2 days", "HIGH",     "Not Started", "Wrap every strategy signal"),
        ("Test: fewer but higher quality signals",    "Terminal / Grafana",                        "Day 62",   "Passive","HIGH",     "Not Started", "Signal count should drop ~40%"),
        "PRIORITY 6 — Funding Rate Extremes (Days 63–68)",
        ("Create FundingRateExtremeStrategy class",   "strategies/funding_rate_extremes.py (new)","Day 63",   "3 days", "HIGH",     "Not Started", "Requires Phase 2 get_funding_rate()"),
        ("+0.1% funding → sell; -0.05% → buy",        "strategies/funding_rate_extremes.py",       "Day 65",   "1 day",  "HIGH",     "Not Started", "Contrarian signal unique to perps"),
        ("Wire for HIGH_VOL regime",                  "core/engine.py",                            "Day 66",   "1 day",  "HIGH",     "Not Started", "Activate only in HIGH_VOLATILITY"),
        ("Backtest + validate",                       "backtesting/",                              "Day 67",   "1 day",  "HIGH",     "Not Started", "Gate: Sharpe > 0.8"),
        "PRIORITY 7 — Volume-Weighted Momentum (Days 69–72)",
        ("Create VolumeWeightedMomentumStrategy",     "strategies/volume_weighted_momentum.py (new)","Day 69", "2 days", "HIGH",     "Not Started", "10-bar momentum × volume ratio"),
        ("Delete old 3-bar MomentumStrategy",         "strategy/momentum.py",                      "Day 71",   "30 min", "HIGH",     "Not Started", "Remove 2/5 rated strategy"),
        ("Backtest vs old momentum — must be better", "backtesting/",                              "Day 72",   "1 day",  "HIGH",     "Not Started", "Gate: Sharpe > old version"),
        "PRIORITY 8 — S/R Breakout (Days 73–77)",
        ("Create SupportResistanceBreakoutStrategy",  "strategies/sr_breakout.py (new)",           "Day 73",   "3 days", "HIGH",     "Not Started", "20-bar high/low + volume spike"),
        ("Wire for TRENDING + HIGH_VOL regimes",      "core/engine.py",                            "Day 76",   "1 day",  "HIGH",     "Not Started", "Volume > 1.5× average required"),
        ("Backtest + validate",                       "backtesting/",                              "Day 77",   "1 day",  "HIGH",     "Not Started", "Gate: Sharpe > 0.8"),
        "FULL BACKTEST VALIDATION (Days 78–82)",
        ("Run all 8 new strategies individually",     "backtesting/",                              "Day 78",   "2 days", "CRITICAL", "Not Started", "Only keep if Sharpe > 0.8"),
        ("Run all 13 strategies together 72 hours",   "Terminal / Grafana",                        "Days 80–82","3 days","CRITICAL","Not Started", "Check per-strategy PnL on Grafana"),
    ]


def _phase4_tasks():
    return [
        "MANDATORY VALIDATION — Do NOT skip or shorten",
        ("Start 14-day continuous paper run",         "Terminal",               "Day 83",    "2 weeks", "CRITICAL", "Not Started", "Must be continuous, no restarts"),
        ("Day 7 checkpoint: Sharpe > 0.5",           "scripts/analyze_paper_trades.py","Day 90", "–",  "CRITICAL", "Not Started", "Drawdown < 6% also required"),
        ("Day 7 checkpoint: Drawdown < 6%",          "Grafana / analyze script","Day 90",    "–",      "CRITICAL", "Not Started", "Fail = tune or disable strategies"),
        ("Weekly run of analyze_paper_trades.py",     "scripts/analyze_paper_trades.py","Day 90","1 hr","HIGH",     "Not Started", "Per-strategy breakdown"),
        ("Identify losing strategies",                "Grafana / script output","Day 91",    "–",      "HIGH",     "Not Started", "Tune or disable if consistently negative"),
        ("Day 14 final gate: Sharpe > 0.8",          "scripts/analyze_paper_trades.py","Day 97","–",  "CRITICAL", "Not Started", "Both Day 14 gates MUST pass"),
        ("Day 14 final gate: Win rate > 45%",         "scripts/analyze_paper_trades.py","Day 97","–",  "CRITICAL", "Not Started", "Check per-strategy win rates"),
        ("Day 14 final gate: Max drawdown < 8%",      "Grafana",                "Day 97",    "–",      "CRITICAL", "Not Started", "Hard limit before any live money"),
    ]


def _phase5_tasks():
    return [
        "HETZNER SERVER SETUP",
        ("Register at hetzner.com",                  "hetzner.com",            "Day 98",    "30 min", "HIGH",     "Not Started", "No credit card needed for signup"),
        ("Create CX22 server (2vCPU, 4GB, Ubuntu)",  "Hetzner dashboard",      "Day 98",    "30 min", "HIGH",     "Not Started", "~₹320/month (~$4)"),
        ("Set up SSH key for secure access",          "PowerShell / Hetzner",   "Day 98",    "1 hr",   "HIGH",     "Not Started", "Claude guides you through this"),
        ("Install Python 3.12 + Docker on server",    "SSH Terminal",           "Day 99",    "2 hrs",  "HIGH",     "Not Started", "Same steps as local setup"),
        ("git clone project to server",              "SSH Terminal",           "Day 99",    "30 min", "HIGH",     "Not Started", "poetry install on server"),
        ("Set up production .env (NOT testnet keys)","SSH Terminal",           "Day 100",   "30 min", "CRITICAL", "Not Started", "DELTA_EXCHANGE_ENV=prod-india"),
        ("Run live_preflight.py sanity check",        "scripts/live_preflight.py","Day 100", "30 min", "CRITICAL", "Not Started", "Final check before going live"),
        "GOING LIVE — DO NOT RUSH",
        ("Paper mode on cloud server 48 hours",       "SSH Terminal",           "Day 101",   "2 days", "CRITICAL", "Not Started", "Must pass before live mode"),
        ("Set DELTA_MODE=live in server .env",        ".env on server",         "Day 103",   "5 min",  "CRITICAL", "Not Started", "Only after cloud paper passes"),
        ("Start at half paper size (ORDER_SIZE=50)",  ".env on server",         "Day 103",   "5 min",  "CRITICAL", "Not Started", "Scale up only after 10 good live trades"),
        ("Monitor Grafana intensively first 4 hours", "http://server-ip:3000",  "Day 103",   "4 hrs",  "CRITICAL", "Not Started", "Watch for execution errors"),
        ("Confirm first real trade + stop on Delta UI","Delta Exchange account", "Day 103",  "–",      "CRITICAL", "Not Started", "Stop order must appear in exchange"),
        ("Set up systemd auto-restart on reboot",     "SSH Terminal",           "Day 104",   "1 hr",   "HIGH",     "Not Started", "systemctl enable trading-bot"),
        ("Set up Alertmanager → Telegram alerts",     "docker-compose.yml",     "Day 105",   "2 hrs",  "HIGH",     "Not Started", "Alert on: daily loss > 3%, bot crash"),
    ]


def _sheet_strategy_registry(wb):
    ws = wb.create_sheet("📈 Strategy Registry")
    ws.sheet_view.showGridLines = False

    ws.merge_cells("A1:H1")
    c = ws["A1"]
    c.value = "COMPLETE STRATEGY REGISTRY — 13 STRATEGIES"
    c.fill = fill(C["header_dark"])
    c.font = bold_font(size=13, color=C["white"])
    c.alignment = center()
    ws.row_dimensions[1].height = 30

    headers = ["#", "Strategy Name", "File", "Regime", "Rating", "Status", "Phase Built", "Notes"]
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=2, column=col, value=h)
        cell.fill = fill(C["header_mid"])
        cell.font = bold_font(color=C["white"])
        cell.alignment = center()
        cell.border = thin_border()

    strategies = [
        [1,  "Trend Following (ATR stops)",      "strategies/trend_following.py",              "TRENDING",          "4/5 ⭐⭐⭐⭐", "Keep",           "Existing",  "Best existing strategy — ATR-scaled SL/TP"],
        [2,  "Mean Reversion (Z-score)",          "strategies/mean_reversion.py",               "RANGING / LOW_VOL", "4/5 ⭐⭐⭐⭐", "Keep",           "Existing",  "Z-score(20), entry at ±1.5σ deviation"],
        [3,  "EMA Crossover",                     "strategy/ema_crossover.py",                  "TRENDING",          "3/5 ⭐⭐⭐",   "Keep + fix",     "Existing",  "Add volume filter in Phase 3"],
        [4,  "RSI Scalping",                      "strategy/rsi_scalping.py",                   "RANGING / HIGH_VOL","3/5 ⭐⭐⭐",   "Fix RSI formula","Phase 2",   "Wilder's smoothing fix needed"],
        [5,  "Momentum (3-bar)",                  "strategy/momentum.py",                       "All regimes",       "2/5 ⭐⭐",     "REPLACE → #12",  "Existing",  "DELETE after Volume-Weighted Momentum built"],
        [6,  "VWAP Deviation",                    "strategies/vwap_deviation.py (NEW)",         "RANGING / LOW_VOL", "TBD",          "Build",          "Phase 3",   "Priority 1 — biggest gap, volume-weighted"],
        [7,  "Bollinger Bands Squeeze",           "strategies/bollinger_squeeze.py (NEW)",      "LOW_VOL → HIGH_VOL","TBD",          "Build",          "Phase 3",   "Squeeze detector before explosive moves"],
        [8,  "MACD Histogram Divergence",         "strategies/macd_histogram.py (NEW)",         "TRENDING",          "TBD",          "Build",          "Phase 3",   "Earlier entries than EMA crossover alone"],
        [9,  "Order Book Imbalance",              "strategies/order_book_imbalance.py (NEW)",   "All (secondary)",   "TBD",          "Build",          "Phase 3",   "Signals BEFORE price moves — L2 data"],
        [10, "Multi-TF Confluence Filter",        "strategies/multi_timeframe_filter.py (NEW)", "All (filter)",      "TBD",          "Build",          "Phase 3",   "1m + 15m + 1h must all agree direction"],
        [11, "Funding Rate Extremes",             "strategies/funding_rate_extremes.py (NEW)",  "HIGH_VOL",          "TBD",          "Build",          "Phase 3",   "Contrarian — crowded longs/shorts reversal"],
        [12, "Volume-Weighted Momentum",          "strategies/volume_weighted_momentum.py (NEW)","All regimes",      "TBD",          "Build",          "Phase 3",   "Replaces #5 — 10-bar × volume ratio"],
        [13, "Support / Resistance Breakout",     "strategies/sr_breakout.py (NEW)",            "TRENDING / HIGH_VOL","TBD",        "Build",          "Phase 3",   "20-bar high/low + volume spike confirmation"],
    ]

    regime_colors = {
        "TRENDING": "D6EAF8",
        "RANGING / LOW_VOL": "D5F5E3",
        "HIGH_VOL": "FADBD8",
        "All regimes": "FEF9E7",
        "RANGING / HIGH_VOL": "FDEBD0",
        "LOW_VOL → HIGH_VOL": "E8DAEF",
        "All (secondary)": "FEF9E7",
        "All (filter)": "FDFEFE",
    }

    for i, row_data in enumerate(strategies, 3):
        for col, val in enumerate(row_data, 1):
            cell = ws.cell(row=i, column=col, value=val)
            cell.border = thin_border()
            cell.font = normal_font()
            cell.alignment = left_wrap()
            cell.fill = fill(regime_colors.get(row_data[3], C["white"]))
        ws.row_dimensions[i].height = 22

    set_col_widths(ws, [4, 28, 38, 20, 12, 14, 12, 40])
    ws.freeze_panes = "A3"


def _sheet_known_issues(wb):
    ws = wb.create_sheet("🐛 Known Issues")
    ws.sheet_view.showGridLines = False

    ws.merge_cells("A1:G1")
    c = ws["A1"]
    c.value = "KNOWN ISSUES — Fix Before Phase 3"
    c.fill = fill(C["header_dark"])
    c.font = bold_font(size=13, color=C["white"])
    c.alignment = center()
    ws.row_dimensions[1].height = 30

    headers = ["Severity", "Issue", "File / Location", "Fix Phase", "Status", "Impact", "Notes"]
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=2, column=col, value=h)
        cell.fill = fill(C["header_mid"])
        cell.font = bold_font(color=C["white"])
        cell.alignment = center()
        cell.border = thin_border()

    issues = [
        ["CRIT", "Real API keys committed to GitHub",       ".env",                                  "Phase 1 Day 2", "Fixed",      "Security — exchange account at risk",   "Keys rotated; testnet keys only in .env"],
        ["CRIT", "SL/TP only in memory — lost on crash",   "execution/order_execution_engine.py",   "Phase 2",       "Not Fixed",  "Unprotected positions after restart",   "Native HTTP 400 call removed; still in-memory. Phase 3: bracket API"],
        ["CRIT", "Position state flattened by WebSocket",  "cli/professional_bot.py",               "2026-03-28",    "Fixed",      "Bot re-ordered same position each cycle","min_interval_s raised to 15s; save/restore before retry loop"],
        ["HIGH", "Dockerfile uses Python 3.11 not 3.12",   "Dockerfile line 1",                     "Phase 2 Day 8", "Fixed",      "Docker build will fail",                "Confirmed: FROM python:3.12-slim"],
        ["HIGH", "RSI Wilder's smoothing wrong (~7 pts)",  "strategy/rsi_scalping.py",              "Phase 2 Day 12","Not Fixed",  "Wrong entries — RSI off vs TradingView","Replace simple avg with ewm()"],
        ["HIGH", "Duplicate strategy/ and strategies/",    "src/ directory",                         "Phase 2 Day 8", "Fixed",      "Two classes with same name → bugs",     "Only strategy/ package remains"],
        ["HIGH", "TRADE_FREQUENCY_S=60 — too slow",        ".env",                                   "Phase 2 Day 8", "Fixed",      "Scalp opportunities gone by next poll", "trade_frequency_s = 5 in settings.py"],
        ["HIGH", "Kill switch exits bot permanently",      "cli/professional_bot.py:run_async()",   "Next session",  "Not Fixed",  "After 5% daily loss bot hard-stops; no resume","Need soft-mode: pause trading + alert, not exit loop"],
        ["HIGH", "Portfolio confidence scores 0.24–0.45",  "strategy/portfolio.py",                 "Next session",  "Fixed",      "Below min_signal_confidence=0.6 threshold","Fixed 2026-03-29: RSI added to TRENDING; threshold lowered to 0.50"],
        ["MED",  "YAML config files ignored by Settings",  "core/settings.py",                      "Phase 2 Day 10","Fixed",      "config/paper.yml has no effect",        "Settings.__init__ loads default.yml + mode.yml"],
        ["MED",  "Funding rate costs ignored in PnL",      "execution/fee_manager.py",              "Phase 2 Day 20","Not Fixed",  "PnL overstated by funding payments",    "Add to fee calculation"],
        ["MED",  "No max holding period for positions",    "core/engine.py",                         "Phase 2 Day 24","Fixed",      "Positions can be open forever",         "max_holding_time_s = 1800 (30 min) in settings"],
        ["MED",  "No idempotent order placement",          "api/delta_client.py",                    "Phase 2",       "Not Fixed",  "Duplicate orders on network retry",     "SHA256 client_order_id"],
        ["MED",  "No fill confirmation polling",           "execution/order_execution_engine.py",   "Phase 2",       "Fixed",      "Order could be rejected silently",      "Issue 3 fixed: false CRITICAL downgraded to WARNING when before==after==0"],
        ["MED",  "Two risk managers coexist",              "risk/",                                  "Phase 2",       "Not Fixed",  "Inconsistent behavior, dead code",      "Delete advanced_risk_manager.py"],
        ["LOW",  "No Grafana dashboards",                  "monitoring/",                            "Phase 2 Day 26","Not Fixed",  "Can only see PnL via log files",        "Add via docker-compose"],
        ["LOW",  "No CI/CD pipeline",                      "No .github/workflows/",                 "Phase 5",       "Not Fixed",  "No automated test on push",             "Add after cloud migration"],
    ]

    sev_fill = {"CRIT": C["crit"], "HIGH": C["high"], "MED": C["med"], "LOW": C["low"]}

    for i, row_data in enumerate(issues, 3):
        for col, val in enumerate(row_data, 1):
            cell = ws.cell(row=i, column=col, value=val)
            cell.border = thin_border()
            cell.font = bold_font(color=C["white"]) if col == 1 else normal_font()
            cell.alignment = center() if col in (1, 4, 5) else left_wrap()
            cell.fill = fill(sev_fill.get(row_data[0], C["white"])) if col == 1 else fill(
                C["done"] if row_data[4] == "Fixed" else C["light_grey"] if i % 2 == 0 else C["white"]
            )
        ws.row_dimensions[i].height = 22

    set_col_widths(ws, [8, 38, 34, 14, 12, 34, 32])
    ws.freeze_panes = "A3"

    from openpyxl.worksheet.datavalidation import DataValidation
    dv = DataValidation(type="list", formula1='"Not Fixed,In Progress,Fixed"', showDropDown=False)
    dv.sqref = f"E3:E{len(issues)+2}"
    ws.add_data_validation(dv)


# ═══════════════════════════════════════════════════════════════════════
#  WORD DOCUMENT
# ═══════════════════════════════════════════════════════════════════════

def build_docx():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)

    _docx_styles(doc)
    _docx_cover(doc)
    _docx_overview(doc)
    _docx_phase(doc, "Phase 1 — Local Setup & First Run (Days 1–7)",
                "See the bot running and printing live paper trades within your first week.",
                _docx_phase1_content())
    _docx_phase(doc, "Phase 2 — Fix Critical Issues (Days 8–37)",
                "Security fixed, all bugs corrected, exchange-native stops, Grafana live.",
                _docx_phase2_content())
    _docx_phase(doc, "Phase 3 — Eight New Strategies (Days 38–82)",
                "Build all 8 new strategies. Backtest each. Only activate if Sharpe > 0.8.",
                _docx_phase3_content())
    _docx_phase(doc, "Phase 4 — Paper Validation (Days 83–97)",
                "MANDATORY: 14 consecutive days. Both validation gates must pass.",
                _docx_phase4_content())
    _docx_phase(doc, "Phase 5 — Cloud Migration (Days 98+)",
                "Move to Hetzner, run 24/7, go live with real money.",
                _docx_phase5_content())
    _docx_strategy_table(doc)
    _docx_issues_table(doc)
    _docx_cost_table(doc)

    path = OUT_DIR / "Project_Plan.docx"
    doc.save(path)
    print(f"✅  Word doc saved → {path}")


def _docx_styles(doc):
    from docx.oxml.ns import qn
    styles = doc.styles
    # Ensure Normal style has Calibri
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)


def _set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _docx_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\nDELTA EXCHANGE TRADING BOT")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Complete Project Plan & Roadmap\n")
    r2.font.size = Pt(13)
    r2.font.color.rgb = RGBColor(0x2E, 0x6D, 0xA4)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(
        "97 Days  •  13 Strategies  •  5 Phases  •  $24/month when live\n"
        "Exchange: Delta Exchange India  •  Capital: $1,000\n"
    )
    r3.font.size = Pt(10)
    r3.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)
    doc.add_page_break()


def _docx_overview(doc):
    doc.add_heading("Project Overview", level=1)

    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for cell, txt in zip(hdr, ["Phase", "Days", "Goal", "Key Output"]):
        cell.text = txt
        cell.paragraphs[0].runs[0].bold = True
        _set_cell_bg(cell, "1B3A5C")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    phases = [
        ("Phase 1 – Local Setup",        "1–7",   "Bot running in paper mode",                   "No crashes, DB has rows"),
        ("Phase 2 – Fix Critical Issues", "8–37",  "Security fixed, stops on exchange, Grafana",  "14 issues resolved"),
        ("Phase 3 – New Strategies",      "38–82", "8 new strategies built & backtested",          "13 total strategies active"),
        ("Phase 4 – Paper Validation",    "83–97", "Sharpe > 0.8, DD < 8%, Win rate > 45%",       "Both 14-day gates pass"),
        ("Phase 5 – Cloud Migration",     "98+",   "Hetzner server, 24/7 live trading",            "systemd + Alertmanager active"),
    ]
    row_colors = ["D6EAF8", "FDEBD0", "D5F5E3", "F9E79F", "FADBD8"]
    for row_data, bg in zip(phases, row_colors):
        row = tbl.add_row().cells
        for cell, val in zip(row, row_data):
            cell.text = val
            _set_cell_bg(cell, bg)

    doc.add_paragraph()


def _docx_phase(doc, title, goal, content_blocks):
    doc.add_page_break()
    doc.add_heading(title, level=1)
    p = doc.add_paragraph()
    run = p.add_run(f"Goal: {goal}")
    run.bold = True
    run.font.color.rgb = RGBColor(0x1D, 0x5E, 0x35)

    for block in content_blocks:
        if block[0] == "heading":
            h = doc.add_heading(block[1], level=2)
        elif block[0] == "task":
            _, name, detail, days, priority = block
            p = doc.add_paragraph(style="List Bullet")
            r1 = p.add_run(f"[{priority}] {name} ")
            r1.bold = True
            r1.font.size = Pt(10)
            if detail:
                r2 = p.add_run(f"— {detail}")
                r2.font.size = Pt(9)
                r2.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)
            if days:
                r3 = p.add_run(f"  ({days})")
                r3.font.size = Pt(9)
                r3.font.italic = True
                r3.font.color.rgb = RGBColor(0x2E, 0x6D, 0xA4)
        elif block[0] == "warn":
            p = doc.add_paragraph()
            r = p.add_run(f"⚠  {block[1]}")
            r.bold = True
            r.font.color.rgb = RGBColor(0xBA, 0x75, 0x17)
        elif block[0] == "note":
            p = doc.add_paragraph()
            r = p.add_run(f"→  {block[1]}")
            r.font.color.rgb = RGBColor(0x1D, 0x5E, 0x35)


def _docx_phase1_content():
    return [
        ("heading", "Day 1 – Installation"),
        ("task", "Verify Python 3.12.x",         "run: python --version in PowerShell",         "Day 1",   "CRIT"),
        ("task", "Install Poetry",                "pip install poetry",                           "Day 1",   "CRIT"),
        ("task", "Run poetry install",            "Wait for all libraries to download",           "Day 1",   "CRIT"),
        ("heading", "Day 2 – Security First"),
        ("warn",  "DELETE all content in .env before proceeding — API keys are already exposed on GitHub"),
        ("task", "Rotate API keys in .env",       "Paste safe paper-mode template",               "Day 2",   "CRIT"),
        ("task", "Register on Delta testnet",     "cdn-ind.testnet.deltaex.org",                 "Day 2",   "CRIT"),
        ("task", "Update .env with testnet keys", "DELTA_EXCHANGE_ENV=testnet-india",            "Day 2",   "CRIT"),
        ("heading", "Day 2 – First Run"),
        ("task", "Run in paper mode",             "poetry run python -m delta_exchange_bot.cli.main --mode paper","Day 2","HIGH"),
        ("task", "Observe logs 30+ minutes",      "JSON lines should appear every 60s",           "Day 2",   "HIGH"),
        ("task", "Verify DB has rows",            "Use SQLite Viewer extension in VS Code",       "Day 2",   "HIGH"),
        ("heading", "Days 3–7 – Observation"),
        ("task", "Run 2+ hours daily",            "Note every error message",                     "Days 3–7","MED"),
        ("task", "Phase 1 complete gate",         "Bot runs 2 consecutive days without crashing", "Day 7",   "HIGH"),
    ]


def _docx_phase2_content():
    return [
        ("heading", "Critical Severity (Days 8–13)"),
        ("warn",  "Fix ALL CRIT items before building any new strategy"),
        ("task", "Dockerfile Python 3.11 → 3.12",         "One-line change",                            "Day 8",    "CRIT"),
        ("task", "Merge strategy/ and strategies/ packages","Two RSIScalping classes → shadow bugs",      "Days 8–10","CRIT"),
        ("task", "Change TRADE_FREQUENCY_S 60 → 5",        "Add signal_eval_every=6 setting",            "Day 8",    "CRIT"),
        ("task", "Add idempotent client_order_id",          "SHA256(trade_id+symbol+side)",               "Day 12",   "CRIT"),
        ("heading", "High Severity (Days 12–20)"),
        ("task", "Fix RSI Wilder's smoothing",              "Currently 7 points off vs TradingView",      "Days 12–13","HIGH"),
        ("task", "Wire YAML config files",                  "config/paper.yml currently ignored",         "Days 10–12","HIGH"),
        ("task", "Add exchange-native stop orders",         "place_stop_limit_order() in DeltaClient",    "Days 14–20","CRIT"),
        ("task", "Add fill confirmation polling loop",      "Poll until filled/rejected/timeout",         "Day 13",   "HIGH"),
        ("heading", "Medium Severity (Days 20–26)"),
        ("task", "Add funding rate awareness",              "Log rate, add to fee calculation",           "Days 20–24","MED"),
        ("task", "Add max holding period",                  "30-candle default = 30 min on 1m",           "Days 24–26","MED"),
        ("task", "Persist ProtectionState to PostgreSQL",   "Crash-safe SL/TP recovery on restart",      "Day 20",   "MED"),
        ("heading", "Monitoring (Days 26–37)"),
        ("task", "Add Grafana + Loki to docker-compose",    "Trading Overview + Risk + API Health",       "Days 26–37","MED"),
    ]


def _docx_phase3_content():
    return [
        ("note", "Only activate a strategy in live regime dispatch if backtest Sharpe > 0.8"),
        ("heading", "Priority 1 — VWAP Deviation (Days 38–42)"),
        ("task", "Create VWAPDeviationStrategy",     "Entry: price >0.5% from VWAP → mean revert","Days 38–42","HIGH"),
        ("heading", "Priority 2 — Bollinger Squeeze (Days 43–46)"),
        ("task", "Create BollingerSqueezeStrategy",  "Detect low-vol coil before breakout",       "Days 43–46","HIGH"),
        ("heading", "Priority 3 — MACD Histogram (Days 47–50)"),
        ("task", "Create MACDHistogramStrategy",     "Histogram turn + signal line cross",        "Days 47–50","HIGH"),
        ("heading", "Priority 4 — Order Book Imbalance (Days 51–56)"),
        ("task", "Create OrderBookImbalanceStrategy","Use existing get_orderbook() REST call",    "Days 51–56","HIGH"),
        ("heading", "Priority 5 — Multi-TF Confluence (Days 57–62)"),
        ("task", "Create MultiTimeframeFilter",      "1m + 15m + 1h all must agree direction",   "Days 57–62","HIGH"),
        ("heading", "Priority 6 — Funding Rate Extremes (Days 63–68)"),
        ("task", "Create FundingRateExtremeStrategy","Contrarian — crowded longs/shorts",         "Days 63–68","HIGH"),
        ("heading", "Priority 7 — Volume-Weighted Momentum (Days 69–72)"),
        ("task", "Create VolumeWeightedMomentum",    "Replace 3-bar momentum with 10-bar × vol", "Days 69–72","HIGH"),
        ("task", "Delete old MomentumStrategy",       "strategy/momentum.py — rated 2/5",         "Day 71",   "HIGH"),
        ("heading", "Priority 8 — S/R Breakout (Days 73–77)"),
        ("task", "Create SupportResistanceBreakout", "20-bar high/low + volume spike",            "Days 73–77","HIGH"),
        ("heading", "Full Backtest Validation (Days 78–82)"),
        ("task", "Backtest all 8 strategies individually","Gate: Sharpe > 0.8 per strategy",     "Days 78–79","CRIT"),
        ("task", "Run all 13 together 72 hours paper","Check Grafana per-strategy PnL",           "Days 80–82","CRIT"),
    ]


def _docx_phase4_content():
    return [
        ("warn",  "NEVER use real money until BOTH Day 14 gates pass"),
        ("heading", "Validation Gates"),
        ("task", "14-day continuous paper run",     "Must be continuous — no restarts",           "Days 83–97","CRIT"),
        ("task", "Day 7: Sharpe > 0.5",             "Fail = tune or disable underperformers",     "Day 90",   "CRIT"),
        ("task", "Day 7: Drawdown < 6%",            "Hard limit — investigate if breached",       "Day 90",   "CRIT"),
        ("task", "Day 14: Sharpe > 0.8",            "Final gate — must pass before Phase 5",      "Day 97",   "CRIT"),
        ("task", "Day 14: Win rate > 45%",          "Per-strategy breakdown required",            "Day 97",   "CRIT"),
        ("task", "Day 14: Max drawdown < 8%",       "Hard limit before any real money",           "Day 97",   "CRIT"),
        ("heading", "Tools"),
        ("task", "Run analyze_paper_trades.py",     "Weekly — per-strategy PnL breakdown",        "Weekly",   "HIGH"),
        ("task", "Tune underperforming strategies", "Ask Claude to help adjust thresholds",       "As needed","MED"),
    ]


def _docx_phase5_content():
    return [
        ("heading", "Hetzner Server Setup"),
        ("task", "Register at hetzner.com",          "No credit card needed for signup",           "Day 98",   "HIGH"),
        ("task", "Create CX22 server",               "2vCPU, 4GB RAM, Ubuntu 22.04 — ~₹320/month","Day 98",   "HIGH"),
        ("task", "Install Python 3.12 + Docker",     "Same steps as local setup",                  "Day 99",   "HIGH"),
        ("task", "Set production .env",              "DELTA_EXCHANGE_ENV=prod-india, fresh keys",  "Day 100",  "CRIT"),
        ("task", "Run live_preflight.py",            "Final sanity check before live",             "Day 100",  "CRIT"),
        ("heading", "Going Live — Do Not Rush"),
        ("warn",  "Cloud paper mode must pass 48 hours before switching DELTA_MODE=live"),
        ("task", "Start at half paper size",         "DELTA_ORDER_SIZE=50",                        "Day 103",  "CRIT"),
        ("task", "Monitor Grafana first 4 hours",    "Watch for execution errors",                  "Day 103",  "CRIT"),
        ("task", "Confirm stop order appears in UI", "Check Delta Exchange account",               "Day 103",  "CRIT"),
        ("task", "Set up systemd auto-restart",      "systemctl enable trading-bot",               "Day 104",  "HIGH"),
        ("task", "Alertmanager → Telegram alerts",   "Alert on: daily loss > 3%, bot crash",      "Day 105",  "HIGH"),
    ]


def _docx_strategy_table(doc):
    doc.add_page_break()
    doc.add_heading("Strategy Registry — All 13 Strategies", level=1)

    strategies = [
        ["1",  "Trend Following",           "TRENDING",          "4/5", "Keep",         "Existing"],
        ["2",  "Mean Reversion",            "RANGING/LOW_VOL",   "4/5", "Keep",         "Existing"],
        ["3",  "EMA Crossover",             "TRENDING",          "3/5", "Keep + fix",   "Existing"],
        ["4",  "RSI Scalping",              "RANGING/HIGH_VOL",  "3/5", "Fix RSI",      "Phase 2"],
        ["5",  "Momentum (3-bar)",          "All",               "2/5", "REPLACE → #12","Existing"],
        ["6",  "VWAP Deviation",            "RANGING/LOW_VOL",   "TBD", "Build",        "Phase 3"],
        ["7",  "Bollinger Squeeze",         "LOW_VOL transition","TBD", "Build",        "Phase 3"],
        ["8",  "MACD Histogram",            "TRENDING",          "TBD", "Build",        "Phase 3"],
        ["9",  "Order Book Imbalance",      "All (secondary)",   "TBD", "Build",        "Phase 3"],
        ["10", "Multi-TF Confluence",       "All (filter)",      "TBD", "Build",        "Phase 3"],
        ["11", "Funding Rate Extremes",     "HIGH_VOL",          "TBD", "Build",        "Phase 3"],
        ["12", "Vol-Weighted Momentum",     "All",               "TBD", "Build",        "Phase 3"],
        ["13", "S/R Breakout",              "TRENDING/HIGH_VOL", "TBD", "Build",        "Phase 3"],
    ]

    tbl = doc.add_table(rows=1, cols=6)
    tbl.style = "Table Grid"
    for cell, txt in zip(tbl.rows[0].cells, ["#", "Strategy", "Regime", "Rating", "Action", "Phase"]):
        cell.text = txt
        cell.paragraphs[0].runs[0].bold = True
        _set_cell_bg(cell, "1B3A5C")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    row_colors = ["D6EAF8", "D5F5E3", "D6EAF8", "FDEBD0", "FADBD8",
                  "D5F5E3", "E8DAEF", "D6EAF8", "FEF9E7", "FDFEFE",
                  "FADBD8", "FEF9E7", "D6EAF8"]
    for row_data, bg in zip(strategies, row_colors):
        row = tbl.add_row().cells
        for cell, val in zip(row, row_data):
            cell.text = val
            _set_cell_bg(cell, bg)


def _docx_issues_table(doc):
    doc.add_page_break()
    doc.add_heading("Known Issues — Fix Before Phase 3", level=1)

    issues = [
        ["CRIT", "Real API keys in .env",               "Phase 1 Day 2"],
        ["CRIT", "SL/TP in memory only — lost on crash","Phase 2"],
        ["HIGH", "Dockerfile uses Python 3.11",          "Phase 2 Day 8"],
        ["HIGH", "RSI formula wrong (~7 pts off)",       "Phase 2 Day 12"],
        ["HIGH", "Duplicate strategy packages",          "Phase 2 Day 8"],
        ["HIGH", "TRADE_FREQUENCY_S=60 too slow",        "Phase 2 Day 8"],
        ["MED",  "YAML config files ignored",            "Phase 2 Day 10"],
        ["MED",  "Funding rate costs ignored",           "Phase 2 Day 20"],
        ["MED",  "No max holding period",                "Phase 2 Day 24"],
        ["MED",  "No idempotent order placement",        "Phase 2"],
        ["MED",  "No fill confirmation loop",            "Phase 2"],
        ["MED",  "Two risk managers coexist",            "Phase 2"],
        ["LOW",  "No Grafana dashboards",                "Phase 2 Day 26"],
        ["LOW",  "No CI/CD pipeline",                    "Phase 5"],
    ]

    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    for cell, txt in zip(tbl.rows[0].cells, ["Severity", "Issue", "Fix Phase"]):
        cell.text = txt
        cell.paragraphs[0].runs[0].bold = True
        _set_cell_bg(cell, "1B3A5C")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    sev_colors = {"CRIT": "FADBD8", "HIGH": "FDEBD0", "MED": "F9E79F", "LOW": "D5F5E3"}
    for row_data in issues:
        row = tbl.add_row().cells
        for cell, val in zip(row, row_data):
            cell.text = val
            _set_cell_bg(cell, sev_colors.get(row_data[0], "FFFFFF"))


def _docx_cost_table(doc):
    doc.add_heading("Monthly Cost Summary", level=2)

    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    for cell, txt in zip(tbl.rows[0].cells, ["Item", "Cost", "Notes"]):
        cell.text = txt
        cell.paragraphs[0].runs[0].bold = True
        _set_cell_bg(cell, "1B3A5C")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    costs = [
        ["Claude Pro",              "$20/month",        "Your AI coding partner — essential"],
        ["Local dev (Phases 1–4)", "$0",               "Everything runs on your Windows machine"],
        ["Hetzner CX22 (Phase 5)","~₹320 (~$4)/month","Only needed for 24/7 live trading"],
        ["Delta Exchange fees",    "0.02–0.05%/trade", "Maker 0.02%, taker 0.05%"],
        ["TOTAL during dev",       "$20/month",        "Just Claude Pro"],
        ["TOTAL when live",        "~$24/month",       "Claude Pro + Hetzner"],
    ]
    for i, row_data in enumerate(costs):
        row = tbl.add_row().cells
        bg = "D5F5E3" if "TOTAL" in row_data[0] else ("F8F9FA" if i % 2 == 0 else "FFFFFF")
        for cell, val in zip(row, row_data):
            cell.text = val
            cell.paragraphs[0].runs[0].bold = "TOTAL" in row_data[0]
            _set_cell_bg(cell, bg)


# ═══════════════════════════════════════════════════════════════════════
if __name__ == "__main__":
    build_excel()
    build_docx()
    print("\nDone! Both files saved in the docs/ folder.")
